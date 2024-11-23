import requests
import os
from docx import Document
from bs4 import BeautifulSoup
from langchain_openai.chat_models.azure import AzureChatOpenAI
from dotenv import load_dotenv
load_dotenv


subscription_key = os.getenv("YOUR_SUBSCRIPTION_KEY")
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "eastus2"
language_destination = "pt-br"

def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apin-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{
        'text': text
    }]
    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }
    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]["translations"][0]["text"]


def translate_documento(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
        translated_text = translator_text(paragraph, language_destination)
        full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    translated_doc.save(path_translated)
    return path_translated


def extract_text_from_url(url):
    response = requests.get(url)

    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        for script_or_style in soup(['script', 'style']):
            script_or_style.decompose()
        texto = soup.get_text(separator= ' ')
        #Limpar texto
        linhas = (line.strip() for line in texto.splitlines())
        parts = (phrase.strip() for line in linhas for phrase in line.split(" "))
        texto_limpo = '\n'.join(part for part in parts if part)
        return texto_limpo
    else:
        print(f"Failed to fetch the URL. Status code: {response.status_code}")
        return None

    text = soup.get_text()
    return text

client = AzureChatOpenAI(
    azure_endpoint = "https://oai-dio-bootcamp-dev-eastus-001.openai.azure.com",
    api_key = os.getenv(""),
    api_version = "2024-02-15-preview",
    deployment_name = "gpt-4o-mini",
    max_retries=0
)


def translate_article(text, lang):
    messages = [
        ("system", "Voce atua como tradutor de textos"),
        ("user", f"Traduza o {text} para o idioma {lang} e responda em markdown")
    ]

    response = client.invoke(messages)
    print(response.content)
    return response.content


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    # Traduz uma string enviada como parametro ao chamar o metodo translator_text
    translator_text("I know you're somewhere out there, somewhere far away", language_destination)

    # Traduz o documento do path enviado
    input_file = ".arquivo-teste.docx"
    translate_documento(input_file)

    # Extrair url
    url = 'https://www.ibm.com/topics/artificial-intelligence'
    textResult = extract_text_from_url(url)
    print(textResult)
    # Traduzindo url extraida
    print(translate_article(textResult, "pt-br"))